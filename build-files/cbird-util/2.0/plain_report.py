#!/usr/bin/env python3
"""
Report generating script for C-BIRD workflow
2024-10-22
@author: Kutluhan Incekara
email: kutluhan.incekara@ct.gov
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
import pandas as pd
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import argparse


def read_amr_report(amr_report):
    df =  pd.read_csv(amr_report, sep = '\t')    
    df2 = df[['Element symbol','Subclass','% Coverage of reference','% Identity to reference', 'Subtype']]
    df2 = df2.rename(columns={'Element symbol':'Gene','Subclass':'AR Subclass/Mechanism','% Coverage of reference':'Coverage(%)','% Identity to reference':'Identity(%)'})
    amr_df = df2.loc[df2['Subtype'] == "AMR"]
    amr_df = amr_df.drop(columns={'Subtype'})
    amr_df = amr_df.sort_values(by=['AR Subclass/Mechanism'], ascending=True)

    point_df = df2.loc[df2['Subtype'] == "POINT"]
    point_df = point_df.drop(columns={'Subtype'})
    point_df = point_df.sort_values(by=['AR Subclass/Mechanism'], ascending=True)

    df_list = [amr_df, point_df]
    return df_list


# Function to add a pandas DataFrame to a docx table
def add_df_to_docx_table(doc, df):
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

    # Add the header rows.
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = str(col)

    # Add the data rows.
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.iat[i, j])

    table.style = 'Light List'

def report(date, labid, taxon, percent, amr_df, point_df, disclaimer, logo1, logo2, line1, line2, line3, line4, line5, line6):
    # Create a new Document
    doc = Document()

    # changing the page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Create a header
    section = doc.sections[0]
    header = section.header

    # Add a table for alignment
    table = header.add_table(rows=1, cols=3, width=Inches(7.5))
    table.alignment = 1  # Center alignment for the table

    # Set column width and add pictures and heading
    left_cell = table.cell(0, 0)
    middle_cell = table.cell(0, 1)
    right_cell = table.cell(0, 2)

    left_cell.width = Inches(1.5)
    middle_cell.width = Inches(5)
    right_cell.width = Inches(1.5)

    # Add left picture
    if logo1 != None:
        left_cell.paragraphs[0].add_run().add_picture(logo1, width=Inches(1.15), height=Inches(1.15))

    # Add heading in the middle
    #heading = middle_cell.paragraphs[0].add_run(htext)
    heading = middle_cell.paragraphs[0]
    heading.add_run(line1)
    heading.add_run("\n")
    heading.add_run(line2)
    heading.add_run("\n")
    heading.add_run(line3)
    heading.add_run("\n")
    heading.add_run(line4)
    heading.add_run("\n")
    heading.add_run(line5)
    heading.add_run("\n")
    heading.add_run(line6)
    middle_cell.paragraphs[0].alignment = 1  # Center alignment for the text
    for run in heading.runs:
        run.bold = True

    # Add right picture
    if logo2 != None:
        right_cell.paragraphs[0].add_run().add_picture(logo2, width=Inches(1.15), height=Inches(1.15))

    # add title and center
    title = doc.add_paragraph('Antimicrobial Resistance Gene Analysis - Whole Genome Sequencing (WGS)')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    # add info
    info = doc.add_paragraph('Specimen Information')
    for run in info.runs:
        run.bold = True

    info_table = doc.add_table(rows=3, cols=2)
    for cell in info_table.columns[0].cells:
        cell.width = Inches(3)
    for row in info_table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Inches(0.2)

    info_table.cell(0,0).text = 'WGS Analysis Date:'
    info_table.cell(0,1).text = date
    info_table.cell(1,0).text = 'DPH Lab ID:'
    info_table.cell(1,1).text = labid
    info_table.cell(2,0).text = 'Patient Name:'

    # add organism
    organism = doc.add_paragraph('\nPredicted Organism')
    organism.runs[0].bold = True

    org_table = doc.add_table(rows=1, cols=2)
    for cell in org_table.columns[0].cells:
        cell.width = Inches(3)
    for row in org_table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Inches(0.2)

    org_table.cell(0,0).text = taxon
    org_table.cell(0,1).text = percent + "%"

    # AMR
    amr = doc.add_paragraph('\nResistance Genes')
    amr.runs[0].bold = True

    add_df_to_docx_table(doc, amr_df)

    point = doc.add_paragraph('\nPoint Mutations')
    point.runs[0].bold = True
    add_df_to_docx_table(doc, point_df)

    # Notes
    notes = doc.add_paragraph('\nNotes')
    notes.runs[0].bold = True
    doc.add_paragraph('Coverage(%) refers to the percentage of the sequenced gene length compared to the reference resistance gene. ', style='List Bullet')
    doc.add_paragraph('Identity(%) refers to the similarity between the identified and the reference genes.', style='List Bullet')

    # Disclaimer
    if disclaimer != None:
        with open(disclaimer) as f:
            txt = f.read()
        note2 = doc.add_paragraph('\nDisclaimer')
        note2.runs[0].bold = True
        doc.add_paragraph(txt)
                         
    # Add page numbers to the footer
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create the page number element
    page_num = OxmlElement('w:fldSimple')
    page_num.set(qn('w:instr'), 'PAGE')
    run = footer.add_run()
    run._r.append(page_num)

    # Add " of " text
    footer.add_run(" of ")

    # Create the total number of pages element
    num_pages = OxmlElement('w:fldSimple')
    num_pages.set(qn('w:instr'), 'NUMPAGES')
    run = footer.add_run()
    run._r.append(num_pages)

    # Save the document
    doc.save(labid +'_report.docx')
    print (labid + '_report.docx was created!')

   # Main #
def main(args):
    date = args.date
    id = args.id
    organism = args.organism
    percent = args.percent
    amr_report = args.amr_report
    disclaimer = args.disclaimer
    logo1 = args.logo1
    logo2 = args.logo2
    line1 = args.line1
    line2 = args.line2
    line3 = args.line3
    line4 = args.line4
    line5 = args.line5
    line6 = args.line6

    dfs = read_amr_report(amr_report)
    report(date, id, organism, percent, dfs[0], dfs[1], disclaimer, logo1, logo2, line1, line2, line3, line4, line5, line6)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Plain report generating script for C-BIRD workflow')

    parser.add_argument('-d', '--date', type=str, help='Analysis date', required=True)
    parser.add_argument('-i', '--id', type=str, help='Lab id', required=True)
    parser.add_argument('-o', '--organism', type=str, help='Predicted organism', required=True)
    parser.add_argument('-p', '--percent',  type=str, help='Predicted organism percent match', required=True)
    parser.add_argument('-a', '--amr_report', type=str, help='Amrfinder report file', required=True) 
    parser.add_argument('-n', '--disclaimer', nargs='?', const=None, type=str, help='disclaimer or additional notes')  
    parser.add_argument('-l', '--logo1',  nargs='?', const=None, type=str, help='header logo left')    
    parser.add_argument('-r', '--logo2',  nargs='?', const=None, type=str, help='header logo right')
    parser.add_argument('-hl1', '--line1', nargs='?', const='', type=str, help='header text line 1') 
    parser.add_argument('-hl2', '--line2', nargs='?', const='', type=str, help='header text line 2') 
    parser.add_argument('-hl3', '--line3', nargs='?', const='', type=str, help='header text line 3')
    parser.add_argument('-hl4', '--line4', nargs='?', const='', type=str, help='header text line 4')
    parser.add_argument('-hl5', '--line5', nargs='?', const='', type=str, help='header text line 5')
    parser.add_argument('-hl6', '--line6', nargs='?', const='', type=str, help='header text line 6')
    args = parser.parse_args()

    main(args)

